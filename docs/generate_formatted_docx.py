from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_PATH = Path("docs/SROMS_Project_Report_Formatted.docx")


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_text = OxmlElement("w:t")
    fld_char_text.text = "Right-click and Update Field to generate TOC"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_text)
    run._r.append(fld_char_end)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.runs[0].font.size = Pt(16)
    elif level == 2:
        h.runs[0].font.size = Pt(14)
    else:
        h.runs[0].font.size = Pt(12)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.line_spacing = 1.3


def add_screenshot_placeholder(doc, title):
    add_heading(doc, title, level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ Insert Screenshot Here ]")
    run.bold = True
    run.font.size = Pt(11)
    box = doc.add_paragraph("_" * 110)
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # Title Page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("SMART RESTAURANT ORDERING MANAGEMENT SYSTEM (SROMS)")
    run.bold = True
    run.font.size = Pt(22)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run("\nA Major Project Report\n").bold = True
    st.add_run("for College Submission").italic = True

    doc.add_paragraph("\n")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Submitted by:\nName: ______________________\nRoll No: ____________________\nClass/Sem: _________________\n")
    info.add_run("\nGuided by:\n___________________________\n")
    info.add_run("\nDepartment of __________________\nCollege Name __________________\nSession: 20__ - 20__")

    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt.add_run(f"\nDate: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_page_break()

    # Certificate + Declaration
    add_heading(doc, "Certificate", level=1)
    add_para(
        doc,
        "This is to certify that the project titled 'Smart Restaurant Ordering Management System (SROMS)' is a bonafide work carried out by the student under my guidance.",
    )
    add_para(
        doc,
        "Guide Signature: ____________________    HOD Signature: ____________________",
    )

    add_heading(doc, "Declaration", level=1)
    add_para(
        doc,
        "I hereby declare that this report is my original work and has not been submitted elsewhere for any degree or diploma.",
    )
    add_para(doc, "Student Signature: ____________________")

    add_heading(doc, "Acknowledgement", level=1)
    add_para(
        doc,
        "I thank my faculty guide, department, friends, and family for their support in completing this project successfully.",
    )

    add_heading(doc, "Abstract", level=1)
    add_para(
        doc,
        "SROMS ek Flask-based web application hai jo restaurant ordering process ko digital banata hai. "
        "Customer QR code scan karke menu dekh sakta hai, cart me item add kar sakta hai, aur direct order place kar sakta hai. "
        "Admin panel se menu, categories, orders aur reports manage hoti hain.",
    )

    doc.add_page_break()

    # TOC
    add_heading(doc, "Table of Contents", level=1)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)
    doc.add_page_break()

    # Chapters
    add_heading(doc, "1. Introduction", level=1)
    add_para(
        doc,
        "Restaurant industry me fast service aur accurate order handling bahut important hai. "
        "Traditional manual process me delay aur order mistakes common hote hain. "
        "SROMS ka objective simple, practical, aur scalable digital ordering solution provide karna hai.",
    )

    add_heading(doc, "2. Problem Statement", level=1)
    add_bullets(
        doc,
        [
            "Manual order lene me zyada time lagta hai.",
            "Peak hours me staff overload hota hai.",
            "Order communication errors generate hote hain.",
            "Real-time status tracking nahi hoti.",
        ],
    )

    add_heading(doc, "3. Proposed System", level=1)
    add_bullets(
        doc,
        [
            "Table-specific QR code based ordering flow.",
            "Customer-friendly digital menu and cart.",
            "Admin dashboard for menu and order management.",
            "SQLite database persistence for records.",
        ],
    )

    add_heading(doc, "4. Objectives", level=1)
    add_bullets(
        doc,
        [
            "Restaurant ordering process automate karna.",
            "Order accuracy improve karna.",
            "Admin side control centralized banana.",
            "System ko simple deployment friendly rakhna.",
        ],
    )

    add_heading(doc, "5. Technology Stack", level=1)
    add_bullets(
        doc,
        [
            "Backend: Flask, SQLAlchemy",
            "Database: SQLite",
            "Utilities: qrcode, Werkzeug",
            "Frontend: HTML templates (Jinja2)",
        ],
    )

    add_heading(doc, "6. System Architecture", level=1)
    add_para(
        doc,
        "Customer mobile browser and admin browser dono Flask app se interact karte hain. "
        "Flask app SQLAlchemy ORM ke through SQLite DB se data read/write karti hai. "
        "QR script table-wise PNG files generate karta hai.",
    )

    add_heading(doc, "7. Database Design", level=1)
    add_bullets(
        doc,
        [
            "tables (table_number, qr_code)",
            "categories (name)",
            "menu_items (name, price, is_available, category_id)",
            "orders (table_id, total_amount, status, created_at)",
            "order_items (order_id, item_id, quantity, price)",
            "admin_users (username, password_hash)",
        ],
    )

    add_heading(doc, "8. Input and Output", level=1)
    add_para(doc, "Input examples:")
    add_bullets(
        doc,
        [
            "URL query: /?table_id=1",
            "Cart form: item_id=3, quantity=2",
            'Order JSON: {"table_id":1,"items":[{"item_id":1,"quantity":2}]}',
        ],
    )
    add_para(doc, "Output examples:")
    add_bullets(
        doc,
        [
            "Menu page rendering with categories and available items",
            "Order success page with order_id and line items",
            "JSON response with status and total_amount",
        ],
    )

    add_heading(doc, "9. Module Explanation (Simple Code Logic)", level=1)
    add_para(
        doc,
        "App Factory app start hone par DB initialize karta hai, blueprints register karta hai, aur seed data load karta hai. "
        "Customer module cart aur checkout manage karta hai; Admin module authentication, menu CRUD, order status aur reports handle karta hai.",
    )

    add_heading(doc, "10. Testing and Results", level=1)
    add_bullets(
        doc,
        [
            "Valid table QR redirects to menu: PASS",
            "Invalid table shows error message: PASS",
            "Cart update and remove operations: PASS",
            "Order API validation checks: PASS",
            "Admin login and protected routes: PASS",
        ],
    )

    add_heading(doc, "11. Security and Validation", level=1)
    add_bullets(
        doc,
        [
            "Password hash stored (plain password store nahi hota).",
            "Admin routes session-based protected hain.",
            "Input validation implemented for quantity/status/IDs.",
            "Unavailable menu items order nahi hote.",
        ],
    )

    add_heading(doc, "12. Conclusion", level=1)
    add_para(
        doc,
        "SROMS successfully restaurant ordering process ko digitize karta hai. "
        "Ye system academic project demo ke liye suitable hai aur future me payment, analytics, aur multi-branch features add karke production-grade banaya ja sakta hai.",
    )

    add_heading(doc, "13. Viva Questions", level=1)
    add_bullets(
        doc,
        [
            "Flask choose karne ka reason kya hai?",
            "Order status lifecycle kaise work karta hai?",
            "Session me table binding kyu important hai?",
            "Password hash ka security benefit kya hai?",
            "Future scale ke liye kaunse changes required honge?",
        ],
    )

    doc.add_page_break()

    add_heading(doc, "Screenshot Placeholders", level=1)
    placeholders = [
        "Screenshot 1: Customer Home (QR Entry)",
        "Screenshot 2: Customer Menu Page",
        "Screenshot 3: Cart Page",
        "Screenshot 4: Order Success Page",
        "Screenshot 5: Admin Login",
        "Screenshot 6: Admin Dashboard",
        "Screenshot 7: Menu Management",
        "Screenshot 8: Orders and Status Update",
        "Screenshot 9: Reports Page",
        "Screenshot 10: QR Codes Folder View",
    ]
    for ph in placeholders:
        add_screenshot_placeholder(doc, ph)

    doc.add_page_break()
    add_heading(doc, "Appendix: Selected Code Snippets", level=1)
    add_para(
        doc,
        "Final submission me aap full code listings `seed.py`, `models.py`, `customer.py`, `admin.py`, aur `generate_qr_codes.py` add kar sakte hain to increase page count to 40-50 pages.",
    )

    add_heading(doc, "References", level=1)
    add_bullets(
        doc,
        [
            "https://flask.palletsprojects.com",
            "https://docs.sqlalchemy.org",
            "https://werkzeug.palletsprojects.com",
            "https://docs.python.org",
            "https://pypi.org/project/qrcode",
        ],
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Generated: {OUT_PATH}")


if __name__ == "__main__":
    main()
